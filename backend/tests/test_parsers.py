"""Unit tests for document parsers."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

import pytest

from backend.models import LINE_KIND, PARAGRAPH_KIND, TABLE_KIND
from backend.routers._headers_common import parse_and_store_headers
from backend.services.parsing import parse_document
from backend.store import headers_path, read_json, upload_objects_path, write_jsonl


def test_parse_txt(tmp_path: Path) -> None:
    sample = tmp_path / "sample.txt"
    sample.write_text("Line one\nLine two\n\nLine three\n", encoding="utf-8")

    objects = parse_document(sample)

    assert len(objects) == 3
    assert all(obj["kind"] == LINE_KIND for obj in objects)
    assert {obj["text"] for obj in objects} == {"Line one", "Line two", "Line three"}


def test_parse_docx(tmp_path: Path) -> None:
    from docx import Document

    sample = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Introduction")
    document.add_paragraph("Details paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Item"
    table.cell(1, 1).text = "42"
    document.save(sample)

    objects = parse_document(sample)

    texts = [obj for obj in objects if obj["kind"] == PARAGRAPH_KIND]
    tables = [obj for obj in objects if obj["kind"] == TABLE_KIND]

    assert any(obj["text"] == "Introduction" for obj in texts)
    assert any("Header" in (obj["text"] or "") for obj in tables)


def test_parse_pdf(tmp_path: Path) -> None:
    import importlib.util

    sample = tmp_path / "sample.pdf"

    if importlib.util.find_spec("fitz") is None:
        pytest.skip("PyMuPDF not installed; skipping PDF generation test")

    import fitz  # type: ignore

    document = fitz.open()
    first_page = document.new_page()
    first_page.insert_text((72, 720), "1 Introduction")
    first_page.insert_text((72, 700), "This is a simple PDF line.")
    second_page = document.new_page()
    second_page.insert_text((72, 720), "2 Requirements")
    second_page.insert_text((72, 700), "Provide two bolts per assembly.")
    document.save(str(sample))
    document.close()

    objects = parse_document(sample)

    assert any(obj["kind"] == LINE_KIND for obj in objects)
    contents = "\n".join(obj["text"] for obj in objects if obj["kind"] == LINE_KIND)
    assert "Introduction" in contents
    assert "Provide two bolts" in contents


def test_parse_and_store_headers_includes_locations(
    tmp_path: Path, monkeypatch
) -> None:
    upload_id = "sample"

    objects = [
        {
            "object_id": "obj-1",
            "file_id": upload_id,
            "kind": "line",
            "text": "1 Introduction",
            "page_index": 0,
            "line_index": 0,
            "order_index": 0,
        },
        {
            "object_id": "obj-2",
            "file_id": upload_id,
            "kind": "line",
            "text": "Project overview and goals.",
            "page_index": 0,
            "line_index": 1,
            "order_index": 1,
        },
        {
            "object_id": "obj-3",
            "file_id": upload_id,
            "kind": "line",
            "text": "2 Scope",
            "page_index": 1,
            "line_index": 0,
            "order_index": 2,
        },
        {
            "object_id": "obj-4",
            "file_id": upload_id,
            "kind": "line",
            "text": "System requirements and constraints.",
            "page_index": 1,
            "line_index": 1,
            "order_index": 3,
        },
    ]

    from backend import store

    monkeypatch.setattr(store, "_TMP_DIR", tmp_path, raising=False)
    tmp_path.mkdir(parents=True, exist_ok=True)

    write_jsonl(upload_objects_path(upload_id), objects)

    response = "\n".join(["#headers#", "1 Introduction", "2 Scope", "#headers#"])
    headers = parse_and_store_headers(upload_id, response)

    assert len(headers) == 2

    intro, scope = headers

    assert intro.page_number == 1
    assert intro.line_number == 1
    assert intro.chunk_text == "1 Introduction\nProject overview and goals."

    assert scope.page_number == 2
    assert scope.line_number == 1
    assert scope.chunk_text == "2 Scope\nSystem requirements and constraints."

    persisted = read_json(headers_path(upload_id))
    assert isinstance(persisted, list)
    assert persisted[0]["page_number"] == 1
    assert persisted[0]["line_number"] == 1
    assert "chunk_text" in persisted[0]
