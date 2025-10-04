"""DOCX parsing utilities."""
from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document

from ...models import PARAGRAPH_KIND, TABLE_KIND


def parse_docx(path: Path) -> list[dict[str, Any]]:
    """Parse a DOCX document into normalized objects."""

    document = Document(path)
    objects: list[dict[str, Any]] = []
    order_index = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        objects.append(
            {
                "object_id": str(uuid4()),
                "file_id": path.stem,
                "kind": PARAGRAPH_KIND,
                "text": text,
                "page_index": None,
                "bbox": None,
                "order_index": order_index,
                "paragraph_index": order_index,
                "metadata": {
                    "source": "docx_parser",
                    "style": paragraph.style.name if paragraph.style else None,
                },
            }
        )
        order_index += 1

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        content_rows = [", ".join(filter(None, row)) for row in rows if any(cell for cell in row)]
        if not content_rows:
            continue
        objects.append(
            {
                "object_id": str(uuid4()),
                "file_id": path.stem,
                "kind": TABLE_KIND,
                "text": "\n".join(content_rows),
                "page_index": None,
                "bbox": None,
                "order_index": order_index,
                "n_rows": len(rows) or None,
                "n_cols": max((len(row) for row in rows), default=0) or None,
                "metadata": {"source": "docx_parser", "rows": rows},
            }
        )
        order_index += 1

    return objects
